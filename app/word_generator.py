import re
import datetime
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls


def generate_word_report(state):
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(6)

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title = doc.add_heading("Исследование трендов", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
        run.font.size = Pt(28)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Аналитический отчёт по технологическим трендам")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = info.add_run(f"Запрос: ")
    r.bold = True
    info.add_run(state.get("query", "—"))

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(f"Сгенерировано: {datetime.datetime.now().strftime('%d.%m.%Y %H:%M')}")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()

    score_text = state.get("score", "")
    if "Оценка устойчивости:" in score_text:
        match = re.search(r"Оценка устойчивости:\s*(\d+(?:\.\d+)?)\s*из\s*10", score_text)
        if not match:
            match = re.search(r"(\d+(?:\.\d+)?)\s*/\s*10", score_text)
        if not match:
            match = re.search(r"(\d+(?:\.\d+)?)\s*из\s*10", score_text)
        if match:
            score_num = float(match.group(1))
            doc.add_heading("Оценка устойчивости тренда", level=1)
            score_p = doc.add_paragraph()
            score_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = score_p.add_run(f"{score_num:.0f} из 10")
            r.font.size = Pt(36)
            r.bold = True
            if score_num >= 7:
                r.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
            elif score_num >= 4:
                r.font.color.rgb = RGBColor(0xF3, 0x9C, 0x12)
            else:
                r.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
            score_text_clean = re.sub(r"Оценка устойчивости:.*?\n", "", score_text).strip()
            _write_markdown_enhanced(doc, score_text_clean)

    doc.add_heading("Глобальный анализ", level=1)
    _write_markdown_enhanced(doc, state.get("global_analysis", ""))
    doc.add_heading("Анализ российского рынка", level=1)
    _write_markdown_enhanced(doc, state.get("russia_analysis", ""))

    full_report = state.get("report", "")
    if full_report:
        doc.add_page_break()
        doc.add_heading("Полный отчёт", level=1)
        _write_markdown_enhanced(doc, full_report)

    return doc


def _set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _set_paragraph_spacing(paragraph, before=0, after=0, line=None):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = Pt(line)


def _add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.style = doc.styles["No Spacing"]
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F5F5F5")
    shading.set(qn("w:val"), "clear")
    p._p.get_or_add_pPr().append(shading)
    _set_paragraph_spacing(p, before=4, after=4)
    return p


def _add_bullet_list(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    _parse_inline_text(p, text)
    return p


def _add_numbered_list(doc, text, level=0):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    _parse_inline_text(p, text)
    return p


def _parse_inline_text(paragraph, text):
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    pattern = r"(\*\*.*?\*\*|\*.*?\*|`[^`]+`|\[([^\]]+)\]\(([^)]+)\))"
    parts = re.split(pattern, text)
    i = 0
    while i < len(parts):
        part = parts[i]
        if not part:
            i += 1
            continue
        if part.startswith("[") and "]" in part and "(" in part:
            if i + 2 < len(parts):
                link_text = parts[i + 1]
                link_url = parts[i + 2]
                if link_url and "example" not in link_url:
                    add_hyperlink(paragraph, link_text, link_url)
                else:
                    r = paragraph.add_run(link_text)
                    r.bold = True
                i += 3
                continue
            else:
                r = paragraph.add_run(part)
                i += 1
                continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
            i += 1
            continue
        if part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            r = paragraph.add_run(part[1:-1])
            r.italic = True
            i += 1
            continue
        if part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0xE0, 0x3E, 0x2D)
            i += 1
            continue
        r = paragraph.add_run(part)
        i += 1


def add_hyperlink(paragraph, text, url):
    if not url or "example" in url:
        paragraph.add_run(f" [{text}]")
        return
    try:
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0563C1")
        rPr.append(color)
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "22")
        rPr.append(sz)
        run.append(rPr)
        text_elem = OxmlElement("w:t")
        text_elem.text = text
        run.append(text_elem)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)
    except Exception:
        paragraph.add_run(f" [{text}]")


def _write_markdown_enhanced(doc, text):
    lines = text.split("\n")
    i = 0
    in_code_block = False
    code_buffer = []
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code_block:
                _add_code_block(doc, "\n".join(code_buffer))
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue
        if not stripped:
            i += 1
            continue
        if stripped.startswith("#"):
            level = len(stripped.split()[0])
            title_text = stripped.lstrip("#").strip()
            h = doc.add_heading(title_text, level=min(level, 4))
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1) if level <= 2 else RGBColor(0x2E, 0x86, 0xDE)
            _set_paragraph_spacing(h, before=12, after=4)
            i += 1
            continue
        if stripped.startswith(">"):
            quote_text = stripped.lstrip(">").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single")
            left.set(qn("w:sz"), "12")
            left.set(qn("w:color"), "2E86DE")
            left.set(qn("w:space"), "8")
            pBdr.append(left)
            pPr.append(pBdr)
            r = p.add_run(quote_text)
            r.italic = True
            r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            _set_paragraph_spacing(p, before=4, after=4)
            i += 1
            continue
        if stripped.replace("-", "").strip() == "" and len(stripped) >= 3:
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:color"), "CCCCCC")
            bottom.set(qn("w:space"), "1")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue
        if "|" in stripped and stripped.startswith("|"):
            table_rows = []
            while i < len(lines) and "|" in lines[i]:
                table_rows.append(lines[i].strip())
                i += 1
            if len(table_rows) >= 2:
                headers = [h.strip() for h in table_rows[0].split("|") if h.strip()]
                data_rows = []
                for row in table_rows[2:]:
                    cells = [c.strip() for c in row.split("|") if c.strip()]
                    if cells:
                        data_rows.append(cells)
                if headers:
                    rows_count = len(data_rows) + 1
                    cols_count = len(headers)
                    table = doc.add_table(rows=rows_count, cols=cols_count)
                    table.style = "Table Grid"
                    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for j, h_text in enumerate(headers):
                        cell = table.rows[0].cells[j]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        r = p.add_run(h_text)
                        r.bold = True
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        r.font.size = Pt(10)
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        _set_cell_shading(cell, "0D47A1")
                    for ri, row in enumerate(data_rows):
                        for cj, cell_text in enumerate(row):
                            if cj < cols_count:
                                cell = table.rows[ri + 1].cells[cj]
                                cell.text = ""
                                p = cell.paragraphs[0]
                                _parse_inline_text(p, cell_text)
                                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                if ri % 2 == 1:
                                    _set_cell_shading(cell, "F2F7FC")
                    doc.add_paragraph()
            continue
        if stripped.startswith("- ") or stripped.startswith("* "):
            leading_spaces = len(line) - len(line.lstrip())
            level = leading_spaces // 2
            text_content = stripped[2:].strip()
            _add_bullet_list(doc, text_content, min(level, 3))
            i += 1
            continue
        if re.match(r"^\d+[\.\)]\s", stripped):
            text_content = re.sub(r"^\d+[\.\)]\s", "", stripped)
            _add_numbered_list(doc, text_content)
            i += 1
            continue
        p = doc.add_paragraph()
        _parse_inline_text(p, stripped)
        _set_paragraph_spacing(p, before=2, after=2)
        i += 1
    if code_buffer:
        _add_code_block(doc, "\n".join(code_buffer))